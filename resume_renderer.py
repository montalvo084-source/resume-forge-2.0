import io
import os
from jinja2 import Environment, FileSystemLoader


def _get_jinja_env():
    templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
    return Environment(loader=FileSystemLoader(templates_dir))


def render_html(parsed_resume: dict, user_name: str) -> str:
    env = _get_jinja_env()
    template = env.get_template('resume_template.html')
    return template.render(resume=parsed_resume, user_name=user_name)


def render_pdf(parsed_resume: dict, user_name: str) -> bytes:
    from fpdf import FPDF

    class ResumePDF(FPDF):
        def section_header(self, title):
            self.set_font('Helvetica', 'B', 9)
            self.set_text_color(30, 30, 30)
            self.ln(3)
            self.cell(0, 5, title.upper(), ln=True)
            self.set_draw_color(180, 180, 180)
            self.line(self.get_x(), self.get_y(), self.get_x() + 190, self.get_y())
            self.ln(2)
            self.set_text_color(30, 30, 30)

    pdf = ResumePDF(format='Letter')
    pdf.set_margins(25, 20, 25)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    contact = parsed_resume.get('contact', {})
    name = contact.get('name', user_name)

    # Name
    pdf.set_font('Helvetica', 'B', 20)
    pdf.set_text_color(20, 20, 20)
    pdf.cell(0, 10, name, ln=True, align='C')

    # Contact line
    contact_parts = [v for k, v in contact.items() if k != 'name' and v]
    if contact_parts:
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 5, '  |  '.join(contact_parts), ln=True, align='C')

    pdf.set_draw_color(30, 30, 30)
    pdf.line(25, pdf.get_y() + 2, 185, pdf.get_y() + 2)
    pdf.ln(5)

    # Summary
    if parsed_resume.get('summary'):
        pdf.section_header('Summary')
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 5, parsed_resume['summary'])
        pdf.ln(2)

    # Experience
    if parsed_resume.get('experience'):
        pdf.section_header('Experience')
        for job in parsed_resume['experience']:
            company = job.get('company') or job.get('title', '')
            dates = job.get('dates', '')
            title = job.get('title', '') if job.get('company') else ''

            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(20, 20, 20)
            # Company left, dates right
            pdf.cell(0, 5, f"{company}  {dates}", ln=True)
            if title:
                pdf.set_font('Helvetica', 'I', 10)
                pdf.set_text_color(60, 60, 60)
                pdf.cell(0, 4, title, ln=True)

            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(30, 30, 30)
            for bullet in job.get('bullets', []):
                pdf.cell(5, 5, '')  # indent
                pdf.multi_cell(0, 5, f'-  {bullet}')
            pdf.ln(2)

    # Skills
    if parsed_resume.get('skills'):
        pdf.section_header('Skills')
        pdf.set_text_color(30, 30, 30)
        for category, items in parsed_resume['skills'].items():
            pdf.set_font('Helvetica', '', 10)
            pdf.multi_cell(0, 5, f"{category}: {', '.join(items)}")
        pdf.ln(2)

    # Education
    if parsed_resume.get('education'):
        pdf.section_header('Education')
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(30, 30, 30)
        for edu in parsed_resume['education']:
            line = edu.get('degree', '')
            if edu.get('school'):
                line += f"  -  {edu['school']}"
            if edu.get('year'):
                line += f"  ({edu['year']})"
            pdf.cell(0, 5, line, ln=True)

    return bytes(pdf.output())


def render_docx(parsed_resume: dict, user_name: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Remove default paragraph spacing
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    contact = parsed_resume.get('contact', {})
    name = contact.get('name', user_name)

    # Name heading
    h = doc.add_heading(name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact_parts = [v for k, v in contact.items() if k != 'name' and v]
    if contact_parts:
        p = doc.add_paragraph(' | '.join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Summary
    if parsed_resume.get('summary'):
        doc.add_heading('SUMMARY', level=2)
        doc.add_paragraph(parsed_resume['summary'])

    # Experience
    if parsed_resume.get('experience'):
        doc.add_heading('EXPERIENCE', level=2)
        for job in parsed_resume['experience']:
            p = doc.add_paragraph()
            company = job.get('company') or job.get('title', '')
            run = p.add_run(company)
            run.bold = True
            if job.get('dates'):
                tab_run = p.add_run(f'\t{job["dates"]}')
                tab_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                tab_run.font.size = Pt(10)
            if job.get('company') and job.get('title'):
                title_p = doc.add_paragraph(job['title'])
                title_p.runs[0].italic = True
                title_p.runs[0].font.size = Pt(10)
            for bullet in job.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet')

    # Skills
    if parsed_resume.get('skills'):
        doc.add_heading('SKILLS', level=2)
        for category, items in parsed_resume['skills'].items():
            p = doc.add_paragraph()
            run = p.add_run(f'{category}: ')
            run.bold = True
            p.add_run(', '.join(items))

    # Education
    if parsed_resume.get('education'):
        doc.add_heading('EDUCATION', level=2)
        for edu in parsed_resume['education']:
            p = doc.add_paragraph()
            degree_run = p.add_run(edu.get('degree', ''))
            degree_run.bold = True
            if edu.get('school'):
                p.add_run(f' \u2014 {edu["school"]}')
            if edu.get('year'):
                year_run = p.add_run(f'  {edu["year"]}')
                year_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
